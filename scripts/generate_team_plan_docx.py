"""Generate the consolidated Pitch Factory team plan as a Word document.

Sources consolidated by this generator:
  * Continuation Development Backlog v1.0 (1 September 2026) - binding engineering plan
  * Sales Pitch Factory Conceptual Framework v0.1 - draft, pending sponsor approval
  * Head of AI direction (3 September 2026) - client pack intake, Borek RAG, Gamma rendering
  * O1-O8 decision proposals prepared for Konstantin Borek / Fiona Oldenburg

Usage:  py -3 scripts/generate_team_plan_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

OUTPUT = Path(__file__).resolve().parents[1] / "Borek_Pitch_Factory_Consolidated_Team_Plan_v1.0.docx"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GREY = RGBColor(0x59, 0x5F, 0x6B)
HEADER_FILL = "1B2A4A"
ALT_FILL = "F2F4F7"


# --------------------------------------------------------------------------- helpers


def shade(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, white: bool = False, size: int = 9) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    for index, chunk in enumerate(str(text).split("\n")):
        target = para if index == 0 else cell.add_paragraph()
        target.paragraph_format.space_before = Pt(0)
        target.paragraph_format.space_after = Pt(0)
        run = target.add_run(chunk)
        run.bold = bold
        run.font.size = Pt(size)
        if white:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, white=True)
        shade(cell, HEADER_FILL)
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if r % 2 == 1:
                shade(cells[i], ALT_FILL)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def h1(doc: Document, text: str) -> None:
    para = doc.add_heading(text, level=1)
    for run in para.runs:
        run.font.color.rgb = NAVY


def h2(doc: Document, text: str) -> None:
    para = doc.add_heading(text, level=2)
    for run in para.runs:
        run.font.color.rgb = NAVY


def h3(doc: Document, text: str) -> None:
    para = doc.add_heading(text, level=3)
    for run in para.runs:
        run.font.color.rgb = NAVY


def para(doc: Document, text: str, *, italic: bool = False, size: int = 10.5) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)


def bullets(doc: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        # Support a leading "Label: rest" pattern with a bold label.
        if ": " in item and item.index(": ") < 60:
            label, rest = item.split(": ", 1)
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(10.5)
            run2 = p.add_run(rest)
            run2.font.size = Pt(10.5)
        else:
            run = p.add_run(item)
            run.font.size = Pt(10.5)


def callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Intense Quote")
    run = p.add_run(text)
    run.font.size = Pt(10.5)


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# --------------------------------------------------------------------------- document


def build() -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    section = doc.sections[0]
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    footer = section.footer.paragraphs[0]
    footer.text = "BOREK / PITCH FACTORY  |  Consolidated Team Plan v1.0  |  Internal working document"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GREY

    # ---------------------------------------------------------------- cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("BOREK  /  PITCH FACTORY")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

    heading = doc.add_paragraph()
    run = heading.add_run("Consolidated Team Plan")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    run = sub.add_run(
        "What is confirmed, what is still open, who does what, and in which phase — "
        "combining the Continuation Development Backlog, the Sales Pitch Factory conceptual "
        "framework, the Head of AI direction and the O1–O8 decision proposals."
    )
    run.font.size = Pt(12)
    run.font.color.rgb = GREY

    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Version", "1.0 — internal working plan"],
            ["Date", "3 September 2026"],
            ["Prepared by", "Arvanit Telaku (platform / delivery)"],
            ["Audience", "Arvanit, Endrit, Blenard, Jaya, Mayank + technical leadership"],
            [
                "Consolidates",
                "Continuation Development Backlog v1.0 (1 Sep 2026)\n"
                "Sales Pitch Factory Conceptual Framework v0.1 (draft, 1 Sep 2026)\n"
                "Head of AI direction (3 Sep 2026)\n"
                "O1–O8 decision proposals (pending sponsor sign-off)",
            ],
            [
                "Status",
                "Sections 1–5 and 7 are agreed engineering plan. Section 6 (new capability) "
                "depends on the decisions in section 9 being signed off.",
            ],
        ],
        widths=[1.4, 5.3],
    )

    callout(
        doc,
        "Bottom line: we do not build a second product. We finish the current release, then extend the "
        "same pipeline with richer intake, a Borek knowledge base (RAG) and Gamma as the presentation "
        "channel. The only true replacement is the final rendering stage."
    )

    page_break(doc)

    # ---------------------------------------------------------------- 1. how to read
    h1(doc, "1. How to read this document")

    para(
        doc,
        "Three different sources describe the Pitch Factory, and they are at different levels of "
        "certainty. Mixing them up is the main risk to the plan, so every item below carries one of "
        "three labels.",
    )

    add_table(
        doc,
        ["Label", "Meaning", "How the team should treat it"],
        [
            [
                "CONFIRMED",
                "Agreed in the Continuation Development Backlog v1.0, or already built and proven in the repository.",
                "Build it. This is the current release commitment.",
            ],
            [
                "DIRECTION",
                "Stated by the Head of AI, or marked [M] (said in the alignment meeting) in the conceptual framework.",
                "Plan for it and prepare the interfaces, but get written confirmation before large investment.",
            ],
            [
                "OPEN",
                "Marked [A] (author proposal) in the conceptual framework, or one of the undecided points O1–O8.",
                "Do not build. Needs a decision from the named owner first.",
            ],
        ],
        widths=[1.0, 2.6, 3.1],
    )

    para(
        doc,
        "The conceptual framework (v0.1) is explicitly a draft for review by Konstantin Borek and "
        "Fiona Oldenburg. It is not yet approved and is not a technical specification. The "
        "Continuation Development Backlog is the only document that is currently binding on "
        "day-to-day development.",
        italic=True,
    )

    page_break(doc)

    # ---------------------------------------------------------------- 2. confirmed today
    h1(doc, "2. What is confirmed today")

    h2(doc, "2.1 The product we are finishing now")

    para(
        doc,
        "The current release is one automated journey. A user signs in, creates a presentation, "
        "uploads the transcript of a client conversation, and the AI builds a 14-chapter Framework — "
        "the structured customer story. A human reviews and approves that Framework. Everything after "
        "approval runs automatically until the presentation is ready to preview and download.",
    )

    add_table(
        doc,
        ["Step", "What happens", "Automatic or human"],
        [
            ["1", "Sign in and create a presentation", "Human"],
            ["2", "Upload transcript(s) of the client conversation", "Human"],
            ["3", "AI extracts knowledge and builds the 14-chapter Framework", "Automatic"],
            ["4", "Review the customer story; export Framework to Word / PDF / HTML if wanted", "Human"],
            ["5", "Click “Approve & build presentation” — the single governance checkpoint", "Human"],
            ["6", "Presentation planning", "Automatic"],
            ["7", "Slide content generation, validation and compression", "Automatic"],
            ["8", "PPTX and PDF rendering plus previews", "Automatic"],
            ["9", "Live progress while the job runs (5–8 minutes is normal)", "Automatic"],
            ["10", "“Your presentation is ready” — preview and download PowerPoint / PDF", "Human"],
        ],
        widths=[0.5, 4.4, 1.8],
    )

    h2(doc, "2.2 Scope decisions that are settled")

    bullets(
        doc,
        [
            "Human approval stays: explicit Framework approval is the default governance checkpoint. Slides are never generated from a raw transcript, only from a human-confirmed Framework.",
            "Automation after approval: planning, slides, validation, rendering and preview continue with no further clicks in the normal successful path.",
            "Framework exports: real PDF, HTML and a genuine Word DOCX containing all 14 chapters — not HTML renamed as .docx.",
            "Reliability: jobs must survive long generation times, page refreshes and reconnects with no false failure, and must not create duplicate jobs.",
            "No internals on screen: normal users never need to see SlideSpec, schemas, Celery, layout IDs, compression internals, UUIDs or renderer terminology.",
            "Tenant isolation: row-level security must be proven with two real authenticated users.",
        ],
    )

    h2(doc, "2.3 Where the build actually stands")

    para(
        doc,
        "This is the honest snapshot as of the last live run, and it is what the phase plan starts from.",
    )

    add_table(
        doc,
        ["Area", "Status", "Owner of the gap"],
        [
            ["Framework generation (live)", "Working, about 6 minutes", "—"],
            ["Framework Word / PDF / HTML export", "Working", "—"],
            ["Presentation planning (live)", "Working", "—"],
            ["Deck generation (live)", "Fails on slide 1: COVER_01 statBadges exceeds the maximum of 3", "Blenard (BT-9 / BT-15)"],
            ["Full validate_all.py gate", "Fails on the TIMELINE_01 golden image comparison", "Jaya (JJ-22)"],
            ["EXECUTIVE_SUMMARY_01 layout", "Registered but still a stub", "Jaya (JJ-23)"],
            ["Automatic chain after approval", "Not implemented — manual Plan and Deck steps remain", "Blenard (BT-25)"],
            ["Summary-first Framework review UI", "Not implemented", "Jaya (JJ-24)"],
            ["Raw errors shown in the deck screens", "Not customer-facing yet", "Mayank (MS-25)"],
            ["Platform tickets AT-8 / 37 / 38 / 40 / 41 / 47 / 53 / 56 / 57", "Implemented and proven; needs merge to main plus proof attached", "Arvanit"],
        ],
        widths=[2.2, 3.4, 1.3],
    )

    page_break(doc)

    # ---------------------------------------------------------------- 3. new direction
    h1(doc, "3. The new direction from the Head of AI")

    para(
        doc,
        "Three additional requirements were given on 3 September 2026. All three are consistent with "
        "the conceptual framework — they are the same ideas expressed in engineering terms — but they "
        "change what we build in the second half of the pipeline.",
    )

    add_table(
        doc,
        ["New requirement", "What it means concretely", "Matching item in the conceptual framework", "Label"],
        [
            [
                "Richer intake",
                "When uploading the transcript, the user can optionally also upload a client logo and additional information about the client, so the output is as personalised as possible.",
                "In scope [M]: automated integration of client logos; use of client-specific information from conversations; specific details such as location requirements flow into documents.",
                "DIRECTION",
            ],
            [
                "Borek knowledge base (RAG)",
                "A retrieval layer over Borek's own information — what we do, our pricing, how many team members we can staff on a project — so the model uses real company facts instead of inventing them.",
                "In scope [M]: connection to the presentation database, global templates and rate cards; development starts on dummy data.",
                "DIRECTION",
            ],
            [
                "Gamma as the renderer",
                "Presentations are generated through Gamma from one Borek-branded template. The branding is fixed; only the content changes per client.",
                "In scope [M]: a design agent enforcing corporate identity from a CI sheet. Appendix B already records “Gamma” as a tool to be confirmed.",
                "DIRECTION",
            ],
        ],
        widths=[1.2, 2.4, 2.5, 0.8],
    )

    h2(doc, "3.1 Does this need a different pipeline?")

    para(
        doc,
        "No. The front half of the pipeline — intake, knowledge extraction, the Framework, human "
        "approval, jobs, security, exports — stays exactly as it is and gets extended. One new stage "
        "is inserted in the middle (retrieval of Borek facts). The back half — our own presentation "
        "planning, SlideSpec schemas, design system and PPTX renderer — is the part that Gamma "
        "replaces, because Gamma then owns layout and branding.",
    )

    add_table(
        doc,
        ["Layer", "Verdict"],
        [
            ["Authentication, opportunities, transcript upload, jobs, progress, RLS", "Keep and extend"],
            ["Optional client logo and client information at upload", "Small extension of existing intake"],
            ["PII redaction, prompt versioning, AI observability", "Keep — and extend to cover Gamma calls"],
            ["Knowledge extraction and the 14-chapter Framework", "Keep — becomes the grounded input for rendering"],
            ["Framework review, evidence, approval, Word / PDF export", "Keep"],
            ["Borek company facts (services, pricing, staffing, references)", "New — this is the RAG layer"],
            ["Presentation planning → SlideSpec → our renderer → PPTX", "Replace with Gamma, once Gamma is confirmed in writing"],
            ["Automatic filing of every generated document", "New — the knowledge base / archive"],
        ],
        widths=[4.4, 2.3],
    )

    callout(
        doc,
        "One sentence for leadership: we are not rebuilding Pitch Factory. We are adding client-pack "
        "intake and a Borek knowledge RAG, and replacing our slide renderer with Gamma, while keeping "
        "the Framework approval checkpoint in the middle."
    )

    page_break(doc)

    # ---------------------------------------------------------------- 4. target workflow
    h1(doc, "4. The target workflow when everything is complete")

    para(doc, "This is the full flow the team is building towards, stage by stage.")

    add_table(
        doc,
        ["Stage", "What the system does", "Who is responsible", "Status"],
        [
            [
                "1. Trigger",
                "A sales meeting or client conversation is completed and creation begins.",
                "Sales",
                "CONFIRMED",
            ],
            [
                "2. Capture the client pack",
                "Transcript is uploaded. Optionally also: client logo, and additional client information (location requirements, constraints, contacts, priorities).",
                "Sales · built by AT + MS",
                "Transcript CONFIRMED; logo and extra info DIRECTION",
            ],
            [
                "3. Protect confidential data",
                "Per-opportunity PII redaction is applied before anything is sent to an external model, according to the data classification.",
                "ES + AT",
                "Redaction CONFIRMED; classification depends on O4",
            ],
            [
                "4. Build the customer story",
                "Knowledge extraction and the 14-chapter Framework, with per-fact evidence and review signals.",
                "ES + AT + JJ",
                "CONFIRMED",
            ],
            [
                "5. Retrieve Borek facts",
                "RAG returns our services, current rate-card pricing, available team size and references, so the output uses real company data. Starts on dummy data.",
                "AT (infrastructure) + ES (grounding)",
                "DIRECTION",
            ],
            [
                "6. Human review and approval",
                "Summary-first review: executive summary, pain points, requirements, outcomes, open questions and blocking issues. Then “Approve & build presentation”.",
                "Sales · built by JJ + ES",
                "CONFIRMED",
            ],
            [
                "7. Generate the presentation",
                "The confirmed Framework plus retrieved facts plus the client logo fill the named content slots of one Borek-branded Gamma template. Branding is fixed.",
                "BT + JJ + AT",
                "DIRECTION — depends on O1 and O5",
            ],
            [
                "8. Live progress",
                "Customer-facing stage names, elapsed time, no invented percentages, no false timeout, survives refresh and reconnect.",
                "BT + MS",
                "CONFIRMED",
            ],
            [
                "9. Presentation ready",
                "Preview loads automatically; download PowerPoint is the primary action and PDF the secondary one.",
                "JJ",
                "CONFIRMED",
            ],
            [
                "10. Delivery",
                "Sales sends the approved document to the client. The system does not send to clients automatically.",
                "Sales",
                "CONFIRMED (out of scope [A]: fully automated delivery)",
            ],
            [
                "11. Automatic filing",
                "Every generated document is filed automatically and enriches the knowledge base for future work.",
                "AT + MS",
                "DIRECTION — blocked on O2",
            ],
        ],
        widths=[1.1, 2.7, 1.4, 1.5],
    )

    para(
        doc,
        "Two return paths exist throughout: a correction loop (review fails → regenerate) and a "
        "feedback loop (filed documents enrich the data layer, so the next job is better informed).",
    )

    page_break(doc)

    # ---------------------------------------------------------------- 5. phases
    h1(doc, "5. Phase plan")

    para(
        doc,
        "Phases are sequenced by dependency, not by calendar. Within a phase, the workstreams run in "
        "parallel. A phase is finished when its exit criteria are demonstrably met — not when the "
        "code exists.",
    )

    add_table(
        doc,
        ["Phase", "Name", "Purpose", "Exit criteria"],
        [
            [
                "Phase 1",
                "Close the current release",
                "Finish and prove what is already committed in the Continuation Backlog. This is the only phase that is fully funded by confirmed requirements.",
                "Gates A–E pass and BT-27 accepts a clean end-to-end run: sign in, upload, Framework, approve, automatic build, preview, download — with no developer intervention.",
            ],
            [
                "Phase 2",
                "Decisions and technical proof",
                "Unblock the new direction. Get O1–O8 decided, confirm Gamma in writing, and prove the two unknown technologies with small spikes before committing the team.",
                "O1–O8 signed off; a Gamma spike renders one branded slide deck from our data through the API; a RAG spike answers a pricing and a staffing question from dummy data with a traceable source.",
            ],
            [
                "Phase 3",
                "Richer intake and Borek knowledge",
                "Make the output genuinely personalised: optional logo and client information at upload, and Borek facts retrieved rather than invented.",
                "A generated Framework contains client-specific detail from the extra information, plus service, pricing and staffing facts that each trace back to a versioned corpus entry. No number appears without a source.",
            ],
            [
                "Phase 4",
                "Gamma as the presentation channel",
                "Replace our rendering stage with Gamma behind a feature flag, keeping the existing renderer as the fallback until Gamma is proven.",
                "“Approve & build presentation” produces a Borek-branded Gamma deck with the client logo, downloadable from the ready screen; failures produce understandable recovery states; the flag can switch back to the internal renderer.",
            ],
            [
                "Phase 5",
                "Filing, CI hardening and rollout",
                "Close the loop: every document is filed automatically, the approved CI sheet is enforced, and sales is trained.",
                "Every generated artifact is filed in the agreed repository with workflow and approval metadata in Pitch Factory; CI v1.0 is enforced; the success metrics from O7 are being measured.",
            ],
        ],
        widths=[0.7, 1.5, 2.3, 2.2],
    )

    callout(
        doc,
        "Phase 1 must not stall waiting for decisions. Phase 2 runs in parallel with Phase 1 at "
        "leadership level: the team keeps closing tickets while the open points are being decided."
    )

    h2(doc, "5.1 Work that is deliberately paused")

    para(
        doc,
        "If Gamma becomes the presentation channel, further investment in our own slide layouts is "
        "wasted. The following work should be finished only to the level needed to pass Phase 1, and "
        "then stopped rather than extended.",
    )

    bullets(
        doc,
        [
            "New SlideSpec layouts and schemas beyond what already exists.",
            "EXECUTIVE_SUMMARY_01 as a full custom renderer (JJ-23) — see the decision request in section 9.",
            "Further golden-image regression coverage for internal layouts.",
            "Deeper design-token work in our renderer, beyond what the Framework PDF and DOCX export needs.",
        ],
    )

    page_break(doc)

    # ---------------------------------------------------------------- 6. assignments
    h1(doc, "6. Who does what")

    para(
        doc,
        "Each workstream below lists the same four things: the role, the tickets to close in Phase 1, "
        "the new work in later phases, and the dependencies in both directions. Ticket numbers above "
        "the existing ranges (AT-58+, ES-38+, BT-28+, JJ-26+, MS-27+) are new and proposed by this "
        "plan; they are not yet in the official backlog.",
    )

    # ---- Arvanit
    h2(doc, "6.1 Arvanit Telaku (AT) — Platform, APIs, data, security and jobs")

    para(
        doc,
        "Arvanit owns the spine everyone else builds on: contracts, the FastAPI platform, the "
        "database and migrations, authentication and row-level security, the job system, "
        "observability, and the Framework rendering API. Arvanit does not own prompting or "
        "individual slide layouts.",
    )

    h3(doc, "Phase 1 — close and prove")

    add_table(
        doc,
        ["Ticket", "What to deliver", "Definition of done"],
        [
            ["AT-8", "No silent clipping in compression or retry", "Oversized content either becomes valid through semantic compression or fails explicitly once the retry budget is spent"],
            ["AT-37", "Prove migrations build the database from zero", "A clean database is created entirely from repository migrations, and re-running is safe"],
            ["AT-38", "Prove RLS tenant isolation", "User B cannot read or modify User A's opportunities, transcripts, Frameworks or presentations through any supported path"],
            ["AT-40", "Complete the opportunity and transcript API contract", "Transcript deletion validates ownership, cleans up storage and returns structured errors"],
            ["AT-41", "Framework API plus PDF, HTML and DOCX rendering", "A real DOCX opens cleanly in Word with all 14 chapters, correct version and status, in English and German"],
            ["AT-47", "True per-fact Framework evidence and safe editing", "A reviewer sees which source supports each individual fact, can edit, save, reload and see the same values with correct evidence"],
            ["AT-53", "Durable AI observability", "After a real job, durable records show provider, model, prompt version, tokens, latency, retries and cost, surviving restarts, without storing confidential prompt bodies"],
            ["AT-56", "Durable active-job reconnection", "Refreshing or reopening an in-progress opportunity resumes monitoring the original job instead of starting a new one"],
            ["AT-57", "Failed-stage retry and resume", "A late-stage failure resumes from an appropriate checkpoint without repeating successful earlier stages"],
        ],
        widths=[0.7, 2.4, 3.6],
    )

    para(
        doc,
        "All nine are implemented and proven locally. The remaining work is administrative and it is "
        "blocking other people: commit and open the pull request to main, attach the proof outputs to "
        "each ticket, and mark validation failures as non-retryable so the UI does not offer a "
        "pointless retry.",
    )

    h3(doc, "Phases 2–5 — new work")

    bullets(
        doc,
        [
            "Gamma integration spike (Phase 2): confirm the API, authentication, template mechanics, rate limits, what content payload it accepts, and whether a per-generation client logo is supported. Deliver a written go or no-go.",
            "Data classification and allow-list (Phase 2, with the Head of AI): implement the Public / Internal / Client Confidential / Restricted classification and a field allow-list that governs what may be sent to any external model, including Gamma. This closes O4 technically.",
            "AT-58 extended intake (Phase 3): API, storage, validation and migrations for the optional client logo (file type, size, dimensions) and the optional client information fields, per opportunity and protected by RLS.",
            "AT-59 RAG infrastructure (Phase 3): the corpus store, ingestion, versioning and a retrieval API. Rate cards must be a versioned, structured source owned by Commercial or Sales Ops — never free text, and never a number the model originates.",
            "AT-60 Gamma stage (Phase 4): the outbound adapter, credential handling, timeouts and retries, a new pipeline stage, artifact storage, plus extending AT-53 observability and AT-57 retry to cover it.",
            "AT-61 automatic filing (Phase 5): archive every generated artifact with workflow, version, provenance and approval metadata, and connect to the agreed enterprise repository once O2 is decided.",
        ],
    )

    h3(doc, "Dependencies")

    add_table(
        doc,
        ["Direction", "Detail"],
        [
            ["Arvanit is blocked by", "O2 (where the knowledge base lives), O3 (who owns the rate cards), O4 (what may leave the company) and written confirmation of Gamma"],
            ["Arvanit blocks Endrit", "AT-53 → ES-32 (prompt versions); AT-59 retrieval API → ES-39 (grounded chapters)"],
            ["Arvanit blocks Blenard", "AT-8 and AT-56 → BT-25 and BT-26; AT-60 → BT-28 (Gamma stage in the orchestration)"],
            ["Arvanit blocks Jaya", "AT-41 and AT-47 → JJ-24 (review UI)"],
            ["Arvanit blocks Mayank", "AT-56 → MS-24; AT-57 → MS-25; AT-58 → MS-27 (intake UI)"],
            ["Internal order", "AT-37 → AT-38, and AT-40 → AT-41 → AT-47"],
        ],
        widths=[1.7, 5.0],
    )

    page_break(doc)

    # ---- Endrit
    h2(doc, "6.2 Endrit Shemsedini (ES) — Framework intelligence, privacy and grounding")

    para(
        doc,
        "Endrit owns everything between the raw transcript and the structured customer story: "
        "transcript processing, the knowledge model, the 14-chapter Framework, the review summary, "
        "and the privacy behaviour of everything sent to a language model. In the new direction, "
        "Endrit also owns grounding: making sure Borek facts come from retrieval and are never "
        "invented.",
    )

    h3(doc, "Phase 1 — close and prove")

    add_table(
        doc,
        ["Ticket", "What to deliver", "Definition of done"],
        [
            [
                "ES-4",
                "Per-opportunity PII redaction configuration",
                "Two opportunities can hold different redaction settings; the LLM input follows the correct setting for each; changing one does not affect the other; tests cover enabled and disabled",
            ],
            [
                "ES-32",
                "Persist prompt versions with production jobs",
                "A completed production job exposes durable records showing the actual prompt version used for every relevant LLM call, still inspectable after a worker restart",
            ],
            [
                "ES-36",
                "Framework review summary model",
                "A concise structured layer (executive summary, key pain points, key requirements, target outcomes, assumptions, open questions, contradictions, evidence warnings, readiness, blocking items) derived only from the Framework, introducing no new claims, in English and German",
            ],
            [
                "ES-37",
                "Framework review attention signals",
                "Deterministic signals — READY_TO_APPROVE, REVIEW_RECOMMENDED, BLOCKING_CONTRADICTION, MISSING_REQUIRED_INFORMATION, WEAK_EVIDENCE — with human-readable reasons, where a good average score can never hide a blocking condition",
            ],
        ],
        widths=[0.7, 2.2, 3.8],
    )

    h3(doc, "Phases 3–4 — new work")

    bullets(
        doc,
        [
            "ES-38 use the client pack (Phase 3): carry the optional client information into extraction and Framework synthesis as structured context — location requirements, constraints, contacts, stated priorities — so the personalisation is real and traceable, not decorative.",
            "ES-39 grounded company facts (Phase 3): populate service, pricing, staffing and reference content from the RAG corpus only. Every number carries a source reference to a corpus version. If retrieval returns nothing, the field becomes an open question rather than a guess.",
            "ES-40 presentation content contract (Phase 4): produce the content payload that fills the named slots of the Gamma template, from the confirmed Framework plus retrieved facts. Content only — no layout, no styling.",
            "Classification enforcement (Phase 3–4, with Arvanit): apply the O4 field allow-list in the prompt path, so Client Confidential and Restricted content is never sent to an external model, including Gamma.",
        ],
    )

    h3(doc, "Dependencies")

    add_table(
        doc,
        ["Direction", "Detail"],
        [
            ["Endrit is blocked by", "AT-53 → ES-32; AT-59 retrieval API → ES-39; O3 (rate-card owner and update cycle) → ES-39; the Gamma template slot definition (JJ-26) → ES-40"],
            ["Endrit blocks Jaya", "ES-36 and ES-37 → JJ-24 (the summary-first review UI cannot be built before the summary and signals exist)"],
            ["Endrit blocks Blenard", "ES-40 → BT-28 (the Gamma stage needs a defined content payload)"],
            ["Internal order", "ES-36 → ES-37; ES-38 before ES-39 is meaningful end to end"],
        ],
        widths=[1.7, 5.0],
    )

    page_break(doc)

    # ---- Blenard
    h2(doc, "6.3 Blenard Tahiraj (BT) — Orchestration, automation and release integration")

    para(
        doc,
        "Blenard owns the flow between stages and the release itself: making the pipeline run "
        "automatically after approval, turning backend states into a credible progress experience, "
        "and owning final end-to-end acceptance. Blenard is also the owner of the live deck failure "
        "that currently blocks a clean run.",
    )

    h3(doc, "Phase 1 — close and prove")

    add_table(
        doc,
        ["Ticket", "What to deliver", "Definition of done"],
        [
            [
                "BT-9 / BT-15",
                "Fix the live COVER_01 failure",
                "Deck generation no longer fails on slide 1; the generator and prompt respect the maximum of three stat badges. This is the current blocker for any clean live run",
            ],
            [
                "BT-25",
                "Automated presentation pipeline orchestration",
                "A single “Approve & build presentation” action chains planning, slide generation, validation, rendering and previews with no user action in the successful path; Plan Preview becomes optional; repeated clicks, refreshes and reconnects create no duplicate jobs; failure stops at the correct stage and preserves earlier work",
            ],
            [
                "BT-26",
                "Live generation progress experience",
                "Every backend state maps to a customer-facing name (Preparing, Reading transcripts, Identifying key information, Building the customer story, Checking evidence, Structuring presentation, Writing slides, Checking slide quality, Building PowerPoint, Preparing preview); elapsed time is shown; no invented percentages; the false four-minute timeout is gone; running and failed states never appear together; connection loss never declares the job failed",
            ],
            [
                "BT-27",
                "Final automated end-to-end integration gate",
                "A clean user completes the whole journey without developer intervention; English passes end to end and German gets at least smoke coverage; no manual Plan or Deck steps; no technical dead ends; security, provenance and traceability remain intact",
            ],
        ],
        widths=[0.9, 2.0, 3.8],
    )

    para(
        doc,
        "BT-27 is the release gate for the whole team. Individual owners still answer for their own "
        "acceptance criteria, but Phase 1 is not finished until Blenard accepts the run.",
    )

    h3(doc, "Phases 4–5 — new work")

    bullets(
        doc,
        [
            "BT-28 Gamma rendering stage (Phase 4): replace the internal render stage with the Gamma call inside the orchestration, behind a feature flag, keeping the existing renderer as a fallback until Gamma is proven in production.",
            "BT-29 progress mapping for the new stages (Phase 4): add customer-facing names for the new steps, such as “Retrieving Borek information” and “Building your presentation”, following the same rules as BT-26.",
            "BT-30 end-to-end gate, second edition (Phase 5): re-accept the full journey with the client pack, retrieval and Gamma in place, including the failure and recovery paths.",
        ],
    )

    h3(doc, "Dependencies")

    add_table(
        doc,
        ["Direction", "Detail"],
        [
            ["Blenard is blocked by", "AT-8 and the existing job APIs → BT-25; AT-56 → BT-26; AT-60 and ES-40 → BT-28"],
            ["BT-27 needs", "BT-25, BT-26, AT-41, AT-56, JJ-24, JJ-25 and the MS-24 / MS-25 work where it is incorporated"],
            ["Blenard blocks Jaya", "BT-25 and BT-26 → JJ-25 (the ready screen depends on the automatic chain and progress states)"],
            ["Blenard blocks Mayank", "BT-26 → MS-26 (polish comes after the structural workflow changes land)"],
            ["Blenard blocks everyone", "BT-27 is the shared release gate; nobody's Phase 1 ticket is closed until it passes"],
        ],
        widths=[1.7, 5.0],
    )

    page_break(doc)

    # ---- Jaya
    h2(doc, "6.4 Jaya Joshi (JJ) — Provenance, Framework review UX and the result experience")

    para(
        doc,
        "Jaya owns traceability in Group B layouts and the two screens where the product is judged: "
        "the Framework review, where a human decides whether the customer story is right, and the "
        "finished-presentation screen. In the new direction, Jaya also owns the Borek Gamma template "
        "and the mapping from Framework content into its slots.",
    )

    h3(doc, "Phase 1 — close and prove")

    add_table(
        doc,
        ["Ticket", "What to deliver", "Definition of done"],
        [
            [
                "JJ-9",
                "Group B field-level provenance",
                "Every populated non-metadata value in the canonical Group B layouts traces to a real Framework chapter; missing provenance, unknown paths and unknown chapter IDs all fail validation deterministically",
            ],
            [
                "JJ-22",
                "Resolve the Group B golden regression",
                "The TIMELINE_01 mismatch is diagnosed properly — renderer bug fixed, or the reference updated only with visual approval — and the full Group B golden suite passes. This currently blocks the whole validation gate",
            ],
            [
                "JJ-23",
                "Implement EXECUTIVE_SUMMARY_01 end to end",
                "The planner can select it, Stage B generates a valid SlideSpec, the renderer renders it and the golden regression passes — with no invented commercial or pricing content. See the decision request in section 9 before starting: Gamma may make this obsolete",
            ],
            [
                "JJ-24",
                "Summary-first Framework review UI",
                "A first-time user can decide whether the Framework is accurate enough to proceed without reading all 14 chapters; blocking issues are prominent and prevent invalid approval; full detail stays available; the wording is customer-facing; the primary action is “Approve & build presentation”",
            ],
            [
                "JJ-25",
                "Presentation ready experience",
                "“Your presentation is ready” replaces the operational deck screen; the preview loads automatically; PowerPoint download is primary and PDF secondary; layout IDs and diagnostics are hidden behind a details control; missing or partial artifacts are handled gracefully",
            ],
        ],
        widths=[0.7, 2.1, 3.9],
    )

    h3(doc, "Phases 2–4 — new work")

    bullets(
        doc,
        [
            "JJ-26 Borek Gamma template (Phase 2 design, Phase 4 delivery): define the single branded template and its named content slots, and document which Framework chapter feeds which slot. Branding is locked in the template; only content varies per client. This is the concrete form of the “design agent driven by a CI sheet”.",
            "JJ-27 client logo placement rules (Phase 4): define where and how an uploaded client logo appears, and what happens when the logo is missing or poor quality.",
            "JJ-28 ready screen for Gamma output (Phase 4): preview and download the Gamma artifact through the same “Your presentation is ready” experience, so users never see which engine produced the deck.",
        ],
    )

    h3(doc, "Dependencies")

    add_table(
        doc,
        ["Direction", "Detail"],
        [
            ["Jaya is blocked by", "ES-36 and ES-37 plus AT-47 → JJ-24; BT-25 and BT-26 → JJ-25; O1 (output format), O5 (CI sheet) and Gamma access → JJ-26"],
            ["Jaya blocks Blenard", "JJ-24 and JJ-25 → BT-27; JJ-22 blocks the full validation gate today"],
            ["Jaya blocks Endrit", "JJ-26 slot definitions → ES-40 (the content payload cannot be shaped before the slots exist)"],
            ["Jaya blocks Mayank", "JJ-24 and JJ-25 → MS-26 (polish follows the structural screens)"],
            ["Decision risk", "JJ-23 is one to two days of work on our internal renderer. If Gamma is confirmed, that renderer stops being the product's design system"],
        ],
        widths=[1.7, 5.0],
    )

    page_break(doc)

    # ---- Mayank
    h2(doc, "6.5 Mayank Somwani (MS) — Continuity, recovery, intake experience and polish")

    para(
        doc,
        "Mayank owns what happens around the main flow: finding and resuming earlier work, "
        "understanding failures, and the overall visual coherence of the product. In the new "
        "direction, Mayank also owns the extended upload screen, because that is where the "
        "personalisation is captured.",
    )

    h3(doc, "Phase 1 — close and prove")

    add_table(
        doc,
        ["Ticket", "What to deliver", "Definition of done"],
        [
            [
                "MS-24",
                "Recent presentations",
                "A returning user finds and resumes recent work without needing a UUID or browser history; each entry shows client, date and a customer-facing status (Draft, Analyzing, Needs review, Building presentation, Ready, Needs attention) with Open, Resume and Download actions; user isolation is respected; there is a clear empty state",
            ],
            [
                "MS-25",
                "User-friendly failure and recovery",
                "Failures appear as CONNECTION_LOST, STILL_RUNNING, RETRYING, INPUT_REQUIRED, VALIDATION_NEEDS_REVIEW or TERMINAL_FAILURE with the correct next action; no stack traces for normal users; technical detail sits behind a details control; only one dominant banner at a time",
            ],
            [
                "MS-26",
                "Responsive and visual cleanup pass",
                "Key workflow screens are coherent from desktop to mobile with no horizontal overflow and one clear primary action per screen; sidebars collapse at tablet width; UUIDs and layout IDs are hidden; long emails and narrow headers are handled. Business flows are not redesigned independently of BT and JJ work",
            ],
        ],
        widths=[0.7, 2.1, 3.9],
    )

    h3(doc, "Phases 3–5 — new work")

    bullets(
        doc,
        [
            "MS-27 extended intake experience (Phase 3): add the optional client logo upload — with preview, format and size validation and clear errors — and the optional client information fields to the upload screen. Both must read as genuinely optional so the fast path stays fast.",
            "MS-28 recovery states for the new stages (Phase 4): understandable states and correct next actions for retrieval failures, Gamma failures, credential problems and a missing or changed template.",
            "MS-29 archive and history view (Phase 5): extend recent presentations into a view over the filed knowledge base, so past collateral can be found and reused.",
        ],
    )

    h3(doc, "Dependencies")

    add_table(
        doc,
        ["Direction", "Detail"],
        [
            ["Mayank is blocked by", "AT-56 → MS-24; AT-45, AT-56 and AT-57 → MS-25; JJ-24, JJ-25, BT-26 and MS-24 → MS-26; AT-58 → MS-27; AT-60 and BT-28 → MS-28; O2 → MS-29"],
            ["Mayank blocks Blenard", "MS-24 and MS-25 → BT-27 where they are incorporated in the release"],
            ["Sequencing rule", "MS-26 is deliberately last in Phase 1. Polishing screens before BT-25, BT-26, JJ-24 and JJ-25 land means doing the work twice"],
        ],
        widths=[1.7, 5.0],
    )

    page_break(doc)

    # ---------------------------------------------------------------- 7. dependency map
    h1(doc, "7. Dependency map")

    h2(doc, "7.1 Critical chains for Phase 1")

    add_table(
        doc,
        ["Chain", "Why it matters"],
        [
            ["AT-37 → AT-38", "Tenant isolation can only be proven on a database built from migrations"],
            ["AT-40 → AT-41 → AT-47", "The API contract, then rendering, then per-fact evidence"],
            ["AT-53 → ES-32", "Prompt versions are persisted into the durable record that AT-53 introduces"],
            ["ES-36 → ES-37 → JJ-24", "The summary comes first, then the attention signals, then the UI that presents them"],
            ["AT-56 → BT-26", "Progress cannot survive reconnects until job reconnection is durable"],
            ["BT-25 + BT-26 + JJ-24 + JJ-25 + AT-41 + AT-56 → BT-27", "The release gate depends on all of them"],
            ["BT-9 / BT-15 and JJ-22", "Both block a clean live run and the full validation gate today, so they come first"],
        ],
        widths=[2.6, 4.1],
    )

    h2(doc, "7.2 Critical chains for the new capability")

    add_table(
        doc,
        ["Chain", "Why it matters"],
        [
            ["O4 decision → AT classification and allow-list → ES prompt path", "We cannot send client data to Gamma or any external model before the classification exists"],
            ["Gamma written confirmation → AT-60 → BT-28 → JJ-28 and MS-28", "Everything about the new rendering path depends on Gamma being decided"],
            ["O3 rate-card owner → AT-59 corpus → ES-39 grounded pricing", "Pricing in generated documents must trace to a valid rate-card version"],
            ["AT-58 intake API → MS-27 intake UI → ES-38 use of the client pack", "Storage first, then capture, then use in generation"],
            ["JJ-26 template slots → ES-40 content contract → BT-28 Gamma stage", "The content payload can only be shaped once the slots are known"],
            ["O2 knowledge base → AT-61 filing → MS-29 archive view", "Automatic filing needs a decided destination"],
        ],
        widths=[3.0, 3.7],
    )

    h2(doc, "7.3 Who is waiting on whom")

    add_table(
        doc,
        ["Person", "Waiting on", "Others waiting on them"],
        [
            ["Arvanit", "Leadership: O2, O3, O4, Gamma confirmation", "Endrit, Blenard, Jaya, Mayank — all of them"],
            ["Endrit", "Arvanit (AT-53, AT-59), leadership (O3), Jaya (JJ-26)", "Jaya (JJ-24), Blenard (BT-28)"],
            ["Blenard", "Arvanit (AT-8, AT-56, AT-60), Endrit (ES-40), Jaya (JJ-24, JJ-25), Mayank (MS-24, MS-25)", "The entire team, through BT-27"],
            ["Jaya", "Endrit (ES-36, ES-37), Arvanit (AT-47), Blenard (BT-25, BT-26), leadership (O1, O5)", "Blenard (BT-27), Endrit (ES-40), Mayank (MS-26)"],
            ["Mayank", "Arvanit (AT-56, AT-57, AT-58), Jaya (JJ-24, JJ-25), Blenard (BT-26)", "Blenard (BT-27)"],
        ],
        widths=[0.8, 3.1, 2.8],
    )

    page_break(doc)

    # ---------------------------------------------------------------- 8. still open
    h1(doc, "8. What is explicitly not confirmed")

    para(
        doc,
        "These items appear in the conceptual framework as proposals, marked [A]. They are not "
        "requirements and nobody should build them until they survive review.",
    )

    add_table(
        doc,
        ["Item", "Why it is open", "Impact if confirmed"],
        [
            [
                "The middle “Deepening” journey stage — a tailored pitch with references and the client logo",
                "Only First contact and Concretisation came from the meeting. The middle stage was added by the author",
                "A third output type and probably a second Gamma template",
            ],
            [
                "The sales review step inside the process flow",
                "Marked [A], and it overlaps with O6 (who signs off). Our current product already treats human approval as mandatory",
                "Confirms the checkpoint we have already built",
            ],
            [
                "The out-of-scope list — no automated delivery to clients, no CRM replacement, no automated price negotiation or binding quotes, no additional languages, no migration of historical client folders",
                "Proposed by the author to prevent scope creep; not yet endorsed",
                "Protects the delivery plan. Worth getting endorsed explicitly",
            ],
            [
                "Phases 3 to 5 of the conceptual framework — real sources, CI hardening, rollout and training",
                "Marked [A]; only Phases 0 to 2 came from the meeting",
                "Aligns with Phases 3 to 5 of this plan",
            ],
            [
                "A named technical project lead role",
                "Marked [A]",
                "Clarifies who arbitrates architecture decisions",
            ],
            [
                "Names and tools recorded in the notes: “Jamie” as the meeting tool, “Gamma” / “Gamma Lee” as the CI reference, the “Borik standard”, “Laida” and “Euron”",
                "Appendix B lists all of these as transcription to be confirmed",
                "Gamma in particular is now load-bearing, so the name must be confirmed formally",
            ],
        ],
        widths=[2.3, 2.3, 2.1],
    )

    page_break(doc)

    # ---------------------------------------------------------------- 9. decisions
    h1(doc, "9. Decisions we need from leadership")

    h2(doc, "9.1 O1–O8 with our proposals")

    para(
        doc,
        "These eight points were left undecided in the alignment meeting. The proposals below are "
        "ours, submitted for sign-off; they are not yet decisions.",
    )

    add_table(
        doc,
        ["No.", "Open point", "Our proposal", "Why", "Decision by"],
        [
            [
                "O1",
                "Target output formats",
                "Information pack as PDF; pitch as PPTX plus PDF; proposal as DOCX plus PDF; in-app preview for all",
                "Use the native format that matches the business job. Supporting every format for every artifact multiplies rendering and QA complexity with no clear value",
                "Konstantin Borek / Fiona Oldenburg",
            ],
            [
                "O2",
                "Location and structure of the knowledge base",
                "The enterprise document repository (SharePoint if that is Borek's standard) for files, plus Pitch Factory for metadata and audit",
                "Do not build a document-management platform inside Pitch Factory. Keep files in the enterprise repository and keep workflow, versions, provenance and approval metadata with us",
                "Technical project lead",
            ],
            [
                "O3",
                "Ownership and update cycle of the rate cards",
                "A versioned, structured rate-card source owned by Commercial or Sales Ops. The AI never originates a price",
                "Pricing is a deterministic commercial fact. The model may explain approved pricing, but it must not invent it. Every number must be traceable to a valid rate-card version",
                "Konstantin Borek",
            ],
            [
                "O4",
                "Handling of confidential client data during generation",
                "A Public / Internal / Client Confidential / Restricted classification plus a field allow-list for LLM use",
                "PII redaction alone is not enough. Client architecture, commercial facts or strategy can be confidential without containing personal data. Explicit classification gives policy control",
                "Konstantin Borek / technical project lead",
            ],
            [
                "O5",
                "Completion date of the CI sheet",
                "Use the current CI as a provisional v0.9; require an approved machine-readable CI v1.0 before production",
                "Development should not stop while the CI is being finalised, but production quality needs a frozen, versioned source for colours, typography, spacing, layout, imagery and logo rules",
                "CI ownership",
            ],
            [
                "O6",
                "Approval workflow — who signs off a generated document",
                "Sales approves normal collateral; pricing or discount exceptions trigger commercial approval; legal exceptions escalate when relevant",
                "Risk-based approval. A heavy approval chain on every document destroys the automation benefit, while exception-based escalation preserves accountability",
                "Fiona Oldenburg",
            ],
            [
                "O7",
                "Success criteria and metrics",
                "Measure speed, effort reduction, first-pass acceptance, commercial correctness, CI compliance, adoption and traceability",
                "We need proof of business value, not just proof that the AI works. Metrics let us measure ROI, safety and adoption after rollout",
                "Konstantin Borek",
            ],
            [
                "O8",
                "Timeline for recruiting the fullstack developers",
                "Continue the prototype with the current team; add one experienced fullstack or platform engineer before real-source production integration; a second person only if the workload requires it",
                "Do not block Phase 2 on hiring. Stabilise interfaces and workflow first, then add production integration capacity where it creates the most value",
                "Konstantin Borek",
            ],
        ],
        widths=[0.35, 1.25, 1.7, 2.2, 1.2],
    )

    h2(doc, "9.2 Additional decisions created by the new direction")

    para(
        doc,
        "The Head of AI direction raises five questions that O1–O8 do not cover. Each one changes who "
        "works on what, so they should be answered before Phase 3 starts.",
    )

    add_table(
        doc,
        ["No.", "Question", "Why it is urgent", "Our recommendation"],
        [
            [
                "D1",
                "Is Gamma the confirmed presentation engine, and does our internal renderer become legacy?",
                "Our SlideSpec schemas, design system, renderer and golden tests exist because we were doing what Gamma would do. JJ-23 alone is one to two days of work that Gamma may make pointless",
                "Confirm Gamma in writing, keep the internal renderer as a fallback for one release, and pause new layout work immediately",
            ],
            [
                "D2",
                "One template or one per journey stage?",
                "The Head of AI described one branded template where only the content changes. The conceptual framework describes three different outputs — information pack, tailored pitch, priced proposal. O1 also proposes three different formats",
                "Start with one template for the pitch, and treat the information pack and the proposal as later templates once the first one is proven",
            ],
            [
                "D3",
                "Does the Framework approval checkpoint remain once Gamma is in place?",
                "Our governance rule is that slides are only ever generated from a human-confirmed Framework. Letting Gamma work from the transcript directly would remove that control. This is the practical form of O6",
                "Keep the checkpoint. It is the cheapest safeguard we have against wrong content reaching a client",
            ],
            [
                "D4",
                "Where does the corporate identity live — our design tokens or the Gamma template?",
                "Two sources of truth for branding guarantees drift. It also determines whether O5's machine-readable CI sheet targets our renderer or Gamma",
                "The Gamma template becomes the source of truth for presentations; our design tokens remain only for the Framework PDF and DOCX",
            ],
            [
                "D5",
                "Are prices shown in generated documents indicative or offer-grade?",
                "The conceptual framework rules out binding quotation logic, but a proposal with rate-card pricing looks binding to a client",
                "Indicative and clearly labelled, with sales owning the commercial offer, until Commercial explicitly approves offer-grade output",
            ],
        ],
        widths=[0.35, 1.7, 2.6, 2.05],
    )

    page_break(doc)

    # ---------------------------------------------------------------- 10. gates
    h1(doc, "10. Definition of done")

    h2(doc, "10.1 Phase 1 release gates")

    add_table(
        doc,
        ["Gate", "Requirement", "Owner"],
        [
            ["A — Original acceptance", "ES-4 and ES-32 closed; AT-8, AT-40, AT-41, AT-47 and AT-53 closed; AT-37 and AT-38 verified with live proof; JJ-9 and JJ-22 closed; EXECUTIVE_SUMMARY_01 supported via JJ-23", "ES, AT, JJ"],
            ["B — Framework completeness", "PII policy works per opportunity; per-fact evidence is trustworthy; human review stays explicit; PDF, HTML and DOCX export all work; the DOCX contains all 14 chapters with meaningful content; prompt and model activity is durably auditable", "ES, AT"],
            ["C — Automated presentation flow", "One click on “Approve & build presentation” drives planning, slide generation, validation and compression, PPTX and PDF rendering, and the preview, all automatically", "BT"],
            ["D — Reliability and security", "No false four-minute timeout; refresh and reconnect preserve active jobs; no accidental duplicate jobs; failed stages retry or resume sensibly; RLS tenant isolation is proven", "AT, BT"],
            ["E — Product UX", "No user needs to understand SlideSpec, schemas, Celery, layout IDs, compression internals, UUIDs or renderer internals; the Framework review is summary-first; completion emphasises preview and Download PowerPoint; recent work and recovery states are understandable", "JJ, MS"],
        ],
        widths=[1.4, 4.5, 0.8],
    )

    h2(doc, "10.2 Additional gates for the new capability")

    add_table(
        doc,
        ["Gate", "Requirement", "Owner"],
        [
            ["F — Personalised intake", "The optional client logo and client information are captured, stored, protected by RLS, and demonstrably influence the generated output", "AT, MS, ES"],
            ["G — Grounded company facts", "Services, pricing and staffing come from the RAG corpus with a traceable source and version; no number appears in a document without a source; missing facts become open questions rather than guesses", "AT, ES"],
            ["H — Gamma output", "One approved Borek template renders a client-specific deck with the client logo; branding is not editable per generation; failures are understandable; the feature flag can revert to the internal renderer", "AT, BT, JJ"],
            ["I — Filing and audit", "Every generated artifact is filed automatically with workflow, version, provenance and approval metadata; Gamma calls appear in the durable AI observability records", "AT"],
            ["J — Measurable value", "The O7 metrics are instrumented and reported: speed, effort reduction, first-pass acceptance, commercial correctness, CI compliance, adoption and traceability", "AT, leadership"],
        ],
        widths=[1.4, 4.5, 0.8],
    )

    h2(doc, "10.3 Closure rule")

    callout(
        doc,
        "No ticket is closed until every acceptance condition and every required proof is satisfied. "
        "Implementation without proof is not done. This rule is carried over unchanged from the "
        "Continuation Development Backlog."
    )

    page_break(doc)

    # ---------------------------------------------------------------- 11. risks
    h1(doc, "11. Risks")

    add_table(
        doc,
        ["No.", "Risk", "Impact", "Mitigation", "Owner"],
        [
            ["R1", "Output looks generic or recognisably AI-generated", "High", "A binding CI sheet, one locked branded template, and visual approval before delivery", "JJ, CI ownership"],
            ["R2", "Outdated or incorrect pricing in generated proposals", "High", "Rate cards as the single pricing source with a validity date and a named owner; the model never originates a price", "ES, Commercial"],
            ["R3", "Insufficient data quality in the client information", "Medium", "Structured context fields at capture and mandatory entries before generation", "MS, ES"],
            ["R4", "Delay caused by recruiting", "Medium", "Phases 1 to 3 continue on dummy data with the current team", "Leadership"],
            ["R5", "Confidential content passed to external models, including Gamma", "High", "Field classification and an allow-list enforced in the prompt path, on top of the existing PII redaction", "AT, ES"],
            ["R6", "Lack of adoption in sales", "Medium", "Involve sales early and keep approval rights with sales", "Fiona Oldenburg"],
            ["R7", "Effort wasted on our internal renderer while Gamma is undecided", "High", "Pause new layout work now; get the D1 decision before Phase 4", "Delivery"],
            ["R8", "Two sources of truth for branding — our design tokens and the Gamma template", "Medium", "Decide D4; keep our tokens only for Framework PDF and DOCX output", "AT, JJ"],
            ["R9", "The RAG layer answers confidently from stale or wrong corpus data", "High", "Versioned corpus, source references on every retrieved fact, and a named owner per source", "AT, ES"],
            ["R10", "Dependency on an external vendor for the core output", "Medium", "Keep the internal renderer behind a feature flag for one release, and store the generated artifacts ourselves", "BT, AT"],
        ],
        widths=[0.35, 1.9, 0.65, 2.6, 1.2],
    )

    # ---------------------------------------------------------------- 12. next steps
    h1(doc, "12. Immediate next steps")

    add_table(
        doc,
        ["No.", "Action", "Owner", "When"],
        [
            ["1", "Merge the completed platform work to main and attach the ticket proofs", "Arvanit", "This week — it is blocking others"],
            ["2", "Fix the live COVER_01 deck failure and the TIMELINE_01 golden regression", "Blenard, Jaya", "This week — both block a clean run"],
            ["3", "Close ES-4, ES-32, ES-36 and ES-37 so the review UI can be built", "Endrit", "Phase 1"],
            ["4", "Build JJ-24 and JJ-25 once the Endrit and Arvanit dependencies land", "Jaya", "Phase 1"],
            ["5", "Deliver MS-24 and MS-25, and hold MS-26 until the workflow screens are final", "Mayank", "Phase 1"],
            ["6", "Complete BT-25 and BT-26, then run and accept BT-27", "Blenard", "End of Phase 1"],
            ["7", "Submit O1–O8 and D1–D5 for decision, and get Gamma confirmed in writing", "Delivery to leadership", "In parallel, now"],
            ["8", "Run the Gamma and RAG spikes and report go or no-go", "Arvanit", "Phase 2"],
            ["9", "Define the Borek Gamma template and its content slots", "Jaya, CI ownership", "Phase 2"],
            ["10", "Name the owner and update cycle for the rate cards and the other corpus sources", "Leadership", "Before Phase 3"],
        ],
        widths=[0.35, 3.5, 1.5, 1.35],
    )

    doc.add_paragraph()
    para(
        doc,
        "End of document. Phase 1 is a commitment. Phases 2 to 5 are a plan that becomes a commitment "
        "as the decisions in section 9 are signed off.",
        italic=True,
    )

    doc.save(OUTPUT)
    print(f"written: {OUTPUT}")


if __name__ == "__main__":
    build()
