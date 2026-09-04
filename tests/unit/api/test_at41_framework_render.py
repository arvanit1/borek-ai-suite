"""AT-41: framework PDF / HTML / DOCX render contract."""

from __future__ import annotations

import subprocess
import uuid
import zipfile
from html import unescape
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.auth import create_test_access_token
from app.config import settings
from app.main import create_app

USER_ID = uuid.UUID("aaaaaaaa-aaaa-4aaa-8aaa-aaaaaaaaaaaa")
DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _client() -> TestClient:
    return TestClient(create_app())


def _headers() -> dict[str, str]:
    token = create_test_access_token(
        user_id=USER_ID,
        email="owner@example.com",
        secret=settings.SUPABASE_JWT_SECRET,
    )
    return {"Authorization": f"Bearer {token}"}


def _create_framework(client: TestClient, *, confirm: bool = False) -> tuple[str, dict]:
    created = client.post(
        "/opportunities",
        headers=_headers(),
        json={
            "client_name": "Acme Corp",
            "opportunity_name": "Invoice Automation",
            "department": "Finance",
        },
    )
    assert created.status_code == 201, created.text
    opportunity_id = created.json()["id"]
    generate = client.post(
        f"/opportunities/{opportunity_id}/framework/generate",
        headers=_headers(),
    )
    assert generate.status_code == 202, generate.text
    framework_id = generate.json()["framework_version_id"]
    if confirm:
        confirmed = client.post(
            f"/opportunities/{opportunity_id}/framework/confirm",
            headers=_headers(),
            json={"framework_version_id": framework_id},
        )
        assert confirmed.status_code == 200, confirmed.text
    latest = client.get(f"/frameworks/{framework_id}", headers=_headers())
    assert latest.status_code == 200, latest.text
    return framework_id, latest.json()


def _chapter_titles(framework: dict) -> list[str]:
    chapters = sorted(
        framework["framework_json"]["chapters"],
        key=lambda chapter: int(chapter["chapter_id"]),
    )
    assert len(chapters) == 14
    return [str(chapter["title"]) for chapter in chapters]


def test_render_pdf_returns_pdf_bytes() -> None:
    client = _client()
    framework_id, _ = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=pdf",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.content.startswith(b"%PDF")


def test_render_pdf_includes_version_and_client_metadata() -> None:
    from io import BytesIO

    from pypdf import PdfReader

    client = _client()
    framework_id, framework = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=pdf",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    blob = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(response.content)).pages
    )
    assert "Borek" in blob
    assert "Acme Corp" in blob
    assert "v1" in blob or "Version" in blob
    assert str(framework.get("version") or framework["framework_json"].get("version") or 1) in blob


def test_render_html_returns_html() -> None:
    client = _client()
    framework_id, framework = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=html",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    assert "text/html" in response.headers["content-type"]
    body = unescape(response.text)
    for title in _chapter_titles(framework):
        assert title in body


def test_render_docx_returns_valid_docx() -> None:
    client = _client()
    framework_id, _ = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=docx",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    assert DOCX_TYPE in response.headers["content-type"]
    document = Document(BytesIO(response.content))
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    assert len(texts) >= 14


def test_render_docx_contains_all_chapters() -> None:
    client = _client()
    framework_id, framework = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=docx",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    document = Document(BytesIO(response.content))
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style and str(paragraph.style.name).startswith("Heading")
    ]
    heading_blob = "\n".join(headings)
    for title in _chapter_titles(framework):
        assert title in heading_blob


def test_render_draft_shows_draft_status() -> None:
    client = _client()
    framework_id, framework = _create_framework(client, confirm=False)
    assert framework["status"] == "draft"
    response = client.get(
        f"/frameworks/{framework_id}/render?format=docx",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    document = Document(BytesIO(response.content))
    blob = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "DRAFT" in blob
    assert "not confirmed" in blob.lower()


def test_render_confirmed_does_not_show_draft() -> None:
    client = _client()
    framework_id, framework = _create_framework(client, confirm=True)
    assert framework["status"] == "confirmed"
    response = client.get(
        f"/frameworks/{framework_id}/render?format=docx",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    document = Document(BytesIO(response.content))
    blob = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "DRAFT" not in blob
    assert "Confirmed" in blob


def test_render_docx_includes_client_metadata_and_evidence() -> None:
    client = _client()
    framework_id, framework = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=docx",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    document = Document(BytesIO(response.content))
    blob = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "Borek" in blob
    assert "Acme Corp" in blob
    assert "Invoice Automation" in blob
    assert "Quality and readiness" in blob
    assert "Assumptions and open items" in blob
    json_blob = str(framework["framework_json"])
    if "source_refs" in json_blob:
        assert "Sources:" in blob or "C1" in blob or "conversation" in blob.lower()


def test_render_html_includes_overview_and_tables() -> None:
    client = _client()
    framework_id, _ = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=html",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    body = unescape(response.text)
    assert "Acme Corp" in body
    assert "Borek" in body
    assert "<table" in body
    assert "Quality and readiness" in body


def test_render_does_not_mutate_stored_framework() -> None:
    client = _client()
    framework_id, before = _create_framework(client)
    client.get(f"/frameworks/{framework_id}/render?format=docx", headers=_headers())
    after = client.get(f"/frameworks/{framework_id}", headers=_headers())
    assert after.status_code == 200
    assert after.json()["status"] == before["status"]
    assert after.json()["framework_json"]["title"] == before["framework_json"]["title"]


def test_render_docx_includes_tables_and_lists() -> None:
    client = _client()
    framework_id, framework = _create_framework(client)
    framework_json = framework["framework_json"]
    opportunity_id = framework["opportunity_id"]
    chapter = framework_json["chapters"][2]
    chapter["body"] = [
        {
            "block": "bullets",
            "items": ["Manual invoice matching", "Exception handling backlog"],
        },
        {
            "block": "table",
            "caption": "KPI baseline",
            "columns": ["Metric", "Current"],
            "rows": [["Auto-match rate", "42%"]],
        },
    ]
    patch = client.patch(
        f"/opportunities/{opportunity_id}/framework",
        headers=_headers(),
        json={"framework_json": framework_json},
    )
    assert patch.status_code == 200, patch.text

    response = client.get(
        f"/frameworks/{framework_id}/render?format=docx",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    document = Document(BytesIO(response.content))
    styles = {paragraph.style.name for paragraph in document.paragraphs if paragraph.text.strip()}
    blob = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "List Bullet" in styles or "Manual invoice matching" in blob
    assert "Manual invoice matching" in blob
    assert "Auto-match rate" in blob
    assert len(document.tables) >= 1


def _docx_blob(document: Document) -> str:
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        + [paragraph.text for section in document.sections for paragraph in section.footer.paragraphs]
    )


def _assert_valid_opc(content: bytes) -> None:
    assert content.startswith(b"PK")
    with zipfile.ZipFile(BytesIO(content)) as archive:
        names = set(archive.namelist())
        assert "[Content_Types].xml" in names
        assert "word/document.xml" in names
        document_xml = archive.read("word/document.xml")
        ElementTree.fromstring(document_xml)
        assert b"w:document" in document_xml


def _word_exe() -> Path | None:
    candidates = [
        Path(r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE"),
        Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE"),
        Path(r"C:\Program Files\Microsoft Office\Office16\WINWORD.EXE"),
    ]
    return next((path for path in candidates if path.is_file()), None)


def test_render_german_framework() -> None:
    client = _client()
    framework_id, framework = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=docx&lang=de",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    assert DOCX_TYPE in response.headers["content-type"]
    _assert_valid_opc(response.content)
    document = Document(BytesIO(response.content))
    blob = _docx_blob(document)
    assert "Kunden-Framework-Bericht" in blob
    assert "ENTWURF" in blob
    assert "nicht bestätigt" in blob
    assert "Version" in blob
    assert "Qualität und Bereitschaft" in blob
    assert "Annahmen und offene Punkte" in blob
    for title in _chapter_titles(framework):
        assert title in blob
    assert str(framework.get("version") or framework["framework_json"].get("version") or 1) in blob


def test_render_docx_is_real_word_package_not_renamed_html() -> None:
    client = _client()
    framework_id, _ = _create_framework(client)
    response = client.get(
        f"/frameworks/{framework_id}/render?format=docx",
        headers=_headers(),
    )
    assert response.status_code == 200, response.text
    _assert_valid_opc(response.content)
    assert b"<html" not in response.content[:200]


@pytest.mark.skipif(_word_exe() is None, reason="Microsoft Word is not installed")
def test_render_docx_opens_in_word_english_and_german(tmp_path: Path) -> None:
    client = _client()
    framework_id, framework = _create_framework(client)
    titles = _chapter_titles(framework)
    script = tmp_path / "open_word.ps1"
    script.write_text(
            """
param([string]$Path, [string]$OutFile)
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
try {
  $doc = $word.Documents.Open($Path, $false, $true)
  $text = $doc.Content.Text
  $paras = $doc.Paragraphs.Count
  $doc.Close($false)
  $payload = "PARAS=$paras`nTEXT_START`n$text"
  [System.IO.File]::WriteAllText($OutFile, $payload, [System.Text.UTF8Encoding]::new($false))
} finally {
  $word.Quit()
  [System.Runtime.Interopservices.Marshal]::ReleaseComObject($word) | Out-Null
}
""".strip(),
            encoding="utf-8",
        )

    for lang, markers in (
        ("en", ("DRAFT", "Customer Framework Report", "Version")),
        ("de", ("ENTWURF", "Kunden-Framework-Bericht", "Version")),
    ):
        response = client.get(
            f"/frameworks/{framework_id}/render?format=docx&lang={lang}",
            headers=_headers(),
        )
        assert response.status_code == 200, response.text
        path = tmp_path / f"framework-{lang}.docx"
        extracted = tmp_path / f"framework-{lang}.txt"
        path.write_bytes(response.content)
        opened = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script),
                "-Path",
                str(path),
                "-OutFile",
                str(extracted),
            ],
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
        )
        assert opened.returncode == 0, opened.stderr or opened.stdout
        assert extracted.is_file(), opened.stderr or opened.stdout
        output = extracted.read_text(encoding="utf-8")
        assert "PARAS=" in output
        para_count = int(output.split("PARAS=", 1)[1].splitlines()[0])
        assert para_count >= 14
        for title in titles:
            assert title in output
        for marker in markers:
            assert marker in output
