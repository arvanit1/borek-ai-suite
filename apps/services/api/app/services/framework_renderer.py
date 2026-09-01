"""Framework PDF / HTML / DOCX renderers (AT-41)."""

from __future__ import annotations

import html
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.shared import Pt, RGBColor

from services.framework.guardrails import strip_citations
from services.framework.rendering.customer_pdf import render_customer_pdf

_LABELS = {
    "en": {
        "report": "Customer Framework Report",
        "draft_banner": "DRAFT — not confirmed",
        "confirmed": "Confirmed",
        "in_review": "In review — not confirmed",
        "quality": "Quality and readiness",
        "open_items": "Assumptions and open items",
        "no_open_items": "No open items recorded.",
        "version": "Version",
        "status": "Status",
        "department": "Department",
        "opportunity": "Opportunity",
        "language": "Language",
        "ai_used": "AI is used for",
        "ai_not_used": "AI is not used for",
        "item": "Item",
        "detail": "Detail",
        "score": "Score",
        "band": "Band",
        "week": "Week",
        "activities": "Activities",
        "description": "Description",
        "type": "Type",
        "owner": "Owner",
        "consequence": "If different",
    },
    "de": {
        "report": "Kunden-Framework-Bericht",
        "draft_banner": "ENTWURF — nicht bestätigt",
        "confirmed": "Bestätigt",
        "in_review": "In Prüfung — nicht bestätigt",
        "quality": "Qualität und Bereitschaft",
        "open_items": "Annahmen und offene Punkte",
        "no_open_items": "Keine offenen Punkte erfasst.",
        "version": "Version",
        "status": "Status",
        "department": "Abteilung",
        "opportunity": "Opportunity",
        "language": "Sprache",
        "ai_used": "KI wird verwendet für",
        "ai_not_used": "KI wird nicht verwendet für",
        "item": "Punkt",
        "detail": "Detail",
        "score": "Wert",
        "band": "Band",
        "week": "Woche",
        "activities": "Aktivitäten",
        "description": "Beschreibung",
        "type": "Typ",
        "owner": "Verantwortlich",
        "consequence": "Falls abweichend",
    },
}

_NAVY = RGBColor(0x0E, 0x1C, 0x36)
_GOLD = RGBColor(0xC4, 0xA3, 0x5A)
_DRAFT_RED = RGBColor(0x8B, 0x2E, 0x2E)


def _as_mapping(framework: Any) -> dict[str, Any]:
    if hasattr(framework, "model_dump"):
        return framework.model_dump(mode="json")
    return dict(framework)


def _labels(language: str) -> dict[str, str]:
    return _LABELS["de"] if str(language).lower().startswith("de") else _LABELS["en"]


def _text(value: Any) -> str:
    return strip_citations(str(value or "")).strip()


def _status_key(framework: dict[str, Any]) -> str:
    return str(framework.get("status") or "draft").strip().lower()


def _is_unconfirmed(framework: dict[str, Any]) -> bool:
    return _status_key(framework) != "confirmed"


def _status_banner(framework: dict[str, Any], labels: dict[str, str]) -> str:
    status = _status_key(framework)
    if status == "confirmed":
        return labels["confirmed"]
    if status == "in_review":
        return labels["in_review"]
    return labels["draft_banner"]


def render_framework_pdf(framework: Any, language: str = "en") -> bytes:
    payload = _as_mapping(framework)
    render = dict(payload.get("render") or {})
    render["allowed"] = True
    payload["render"] = render
    return render_customer_pdf(payload, lang="de" if str(language).lower().startswith("de") else "en")


def render_framework_html(framework: Any, language: str = "en") -> str:
    payload = _as_mapping(framework)
    labels = _labels(language)
    chapters = _ordered_chapters(payload)
    parts = [
        "<!DOCTYPE html>",
        '<html lang="%s">' % html.escape(language),
        "<head><meta charset='utf-8'>",
        f"<title>{html.escape(_text(payload.get('title') or labels['report']))}</title>",
        "<style>body{font-family:Georgia,serif;max-width:880px;margin:40px auto;color:#1c1916;}"
        "h1{color:#0E1C36;} table{border-collapse:collapse;width:100%;margin:12px 0;}"
        "th,td{border:1px solid #d9d2c5;padding:6px 8px;text-align:left;}"
        "th{background:#e8eef5;} .draft{color:#8B2E2E;font-weight:bold;}"
        ".muted{color:#5c574f;font-style:italic;font-size:0.9em;}</style>",
        "</head><body>",
        f"<p class='muted'>{html.escape(labels['report'])}</p>",
        f"<h1>{html.escape(_text(payload.get('title') or labels['report']))}</h1>",
    ]
    banner = _status_banner(payload, labels)
    banner_class = " draft" if _is_unconfirmed(payload) else ""
    parts.append(f"<p class='muted{banner_class}'>{html.escape(banner)}</p>")
    parts.append(_html_quality(payload, labels))
    for chapter in chapters:
        title = _text(chapter.get("title"))
        cid = _text(chapter.get("chapter_id"))
        parts.append(f"<h2>{html.escape(f'{cid} {title}'.strip())}</h2>")
        parts.extend(_html_body(chapter.get("body")))
    parts.append(f"<h2>{html.escape(labels['open_items'])}</h2>")
    parts.append(_html_open_items(payload, labels))
    parts.append("</body></html>")
    return "\n".join(parts)


def render_framework_docx(
    framework: Any,
    language: str = "en",
    *,
    include_source_refs: bool = False,
) -> bytes:
    """Render a FrameworkObject to a professional Word document with all 14 chapters."""
    payload = _as_mapping(framework)
    labels = _labels(language)
    doc = Document()
    _style_document(doc)
    _add_cover_page(doc, payload, labels, language)
    _add_quality_summary(doc, payload, labels)
    chapters = _ordered_chapters(payload)
    for index, chapter in enumerate(chapters):
        _add_chapter(doc, chapter, include_source_refs=include_source_refs)
        if index < len(chapters) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _add_open_items(doc, payload, labels)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _style_document(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor(0x1C, 0x19, 0x16)
    styles["Heading 1"].font.color.rgb = _NAVY
    styles["Heading 1"].font.size = Pt(16)
    styles["Title"].font.color.rgb = _NAVY


def _add_cover_page(
    doc: Document,
    framework: dict[str, Any],
    labels: dict[str, str],
    language: str,
) -> None:
    kicker = doc.add_paragraph(labels["report"])
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kicker.runs[0].italic = True
    kicker.runs[0].font.color.rgb = _GOLD

    title = doc.add_heading(_text(framework.get("title") or labels["report"]), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    banner = doc.add_paragraph(_status_banner(framework, labels))
    banner.runs[0].bold = True
    if _is_unconfirmed(framework):
        banner.runs[0].font.color.rgb = _DRAFT_RED
        banner.runs[0].font.size = Pt(14)
    else:
        banner.runs[0].font.color.rgb = _NAVY

    meta = [
        (labels["opportunity"], _text(framework.get("opportunity_id"))),
        (labels["department"], _text(framework.get("department"))),
        (labels["status"], _status_banner(framework, labels)),
        (labels["version"], str(framework.get("version") or 1)),
        (labels["language"], language.upper()),
    ]
    _add_table(doc, [labels["item"], labels["detail"]], meta)


def _add_quality_summary(doc: Document, framework: dict[str, Any], labels: dict[str, str]) -> None:
    doc.add_heading(labels["quality"], level=1)
    scores = framework.get("quality_scores") or {}
    rationale = scores.get("rationale") or {}
    rows = []
    for key in ("opportunity_rating", "conversation_quality", "build_readiness"):
        if key in scores:
            rows.append((key.replace("_", " ").title(), f"{scores.get(key)}/100", _text(rationale.get(key))))
    if rows:
        _add_table(doc, [labels["item"], labels["score"], labels["detail"]], rows)
    for kpi in framework.get("kpis") or []:
        if isinstance(kpi, dict):
            doc.add_paragraph(
                f"{_text(kpi.get('name'))}: {_text(kpi.get('baseline'))} → {_text(kpi.get('target'))}",
                style="List Bullet",
            )


def _add_chapter(doc: Document, chapter: dict[str, Any], *, include_source_refs: bool) -> None:
    cid = _text(chapter.get("chapter_id"))
    title = _text(chapter.get("title"))
    doc.add_heading(f"{cid} {title}".strip(), level=1)
    _add_body(doc, chapter.get("body"))
    if include_source_refs:
        refs = chapter.get("source_refs") or []
        if refs:
            note = doc.add_paragraph(_format_source_refs(refs))
            if note.runs:
                note.runs[0].italic = True
                note.runs[0].font.size = Pt(8)


def _add_body(doc: Document, body: Any) -> None:
    if isinstance(body, str):
        if body.strip():
            doc.add_paragraph(_text(body))
        return
    if not isinstance(body, list):
        return
    for block in body:
        if isinstance(block, str) and block.strip():
            doc.add_paragraph(_text(block))
        elif isinstance(block, dict):
            _add_block(doc, block)


def _add_block(doc: Document, block: dict[str, Any]) -> None:
    kind = str(block.get("block") or "")
    if kind == "prose":
        text = _text(block.get("text"))
        if text:
            doc.add_paragraph(text)
        return
    if kind == "bullets":
        for item in block.get("items") or []:
            doc.add_paragraph(_text(item), style="List Bullet")
        return
    if kind == "callout":
        label = _text(block.get("kind") or "Note").title()
        paragraph = doc.add_paragraph(f"{label}. {_text(block.get('text'))}")
        if paragraph.runs:
            paragraph.runs[0].italic = True
        return
    if kind == "kv_rows":
        if block.get("caption"):
            doc.add_paragraph(_text(block.get("caption"))).runs[0].bold = True
        rows = [(_text(row.get("label")), _text(row.get("value"))) for row in block.get("rows") or [] if isinstance(row, dict)]
        if rows:
            _add_table(doc, ["Item", "Detail"], rows)
        return
    if kind == "table":
        if block.get("caption"):
            doc.add_paragraph(_text(block.get("caption"))).runs[0].bold = True
        columns = [_text(col) or " " for col in block.get("columns") or []]
        data_rows = [
            [_text(cell) for cell in (row if isinstance(row, list) else [row])]
            for row in block.get("rows") or []
        ]
        if columns:
            _add_table(doc, columns, data_rows)
        return
    if kind == "process_flow":
        if block.get("caption"):
            doc.add_paragraph(_text(block.get("caption"))).runs[0].bold = True
        nodes = block.get("nodes") or []
        for index, node in enumerate(nodes, start=1):
            label = _text(node.get("label") if isinstance(node, dict) else node)
            kind_label = _text(node.get("kind")) if isinstance(node, dict) else ""
            suffix = f" ({kind_label})" if kind_label else ""
            doc.add_paragraph(f"{index}. {label}{suffix}", style="List Number")
        return
    if kind == "score_bars":
        rows = []
        for item in block.get("items") or []:
            if not isinstance(item, dict):
                continue
            rows.append(
                (
                    _text(item.get("name")),
                    str(item.get("score") or ""),
                    _text(item.get("band")),
                    _text(item.get("explanation")),
                )
            )
        if rows:
            _add_table(doc, ["Name", "Score", "Band", "Detail"], rows)
        return
    if kind == "timeline":
        rows = []
        for week in block.get("weeks") or []:
            if not isinstance(week, dict):
                continue
            items = "; ".join(_text(item) for item in week.get("items") or [])
            rows.append((_text(week.get("id") or week.get("label")), items))
        if rows:
            _add_table(doc, ["Week", "Activities"], rows)
        return
    if kind == "ai_split":
        used = "\n".join(f"• {_text(item)}" for item in block.get("used_for") or [])
        not_used = "\n".join(f"• {_text(item)}" for item in block.get("not_used_for") or [])
        _add_table(doc, ["AI is used for", "AI is not used for"], [(used, not_used)])
        return
    if kind == "sensitivity":
        rows = [(_text(row.get("label")), _text(row.get("detail"))) for row in block.get("rows") or [] if isinstance(row, dict)]
        if rows:
            _add_table(doc, ["Item", "Detail"], rows)
        return
    if kind == "glossary":
        rows = [(_text(term.get("term")), _text(term.get("meaning"))) for term in block.get("terms") or [] if isinstance(term, dict)]
        if rows:
            _add_table(doc, ["Term", "Meaning"], rows)
        return
    _add_untyped_block(doc, block)


def _add_untyped_block(doc: Document, block: dict[str, Any]) -> None:
    skip = {"block", "source_refs", "id"}
    pairs = [(_humanize(key), _text(value)) for key, value in block.items() if key not in skip and _text(value)]
    if not pairs:
        return
    if len(pairs) == 1:
        doc.add_paragraph(f"{pairs[0][0]}: {pairs[0][1]}" if pairs[0][0] else pairs[0][1])
        return
    _add_table(doc, ["Item", "Detail"], pairs)


def _add_open_items(doc: Document, framework: dict[str, Any], labels: dict[str, str]) -> None:
    doc.add_heading(labels["open_items"], level=1)
    items = [item for item in (framework.get("open_items") or []) if isinstance(item, dict)]
    if not items:
        doc.add_paragraph(labels["no_open_items"])
        return
    rows = [
        (
            _text(item.get("description")),
            _text(item.get("item_type")),
            _text(item.get("owner")),
            _text(item.get("consequence_if_different")),
        )
        for item in items
    ]
    _add_table(
        doc,
        [labels["description"], labels["type"], labels["owner"], labels["consequence"]],
        rows,
    )


def _add_table(doc: Document, headers: list[str], rows: list[tuple[Any, ...] | list[Any]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=max(len(headers), 1))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            if paragraph.runs:
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.color.rgb = _NAVY
        _shade_cell(cell, "E8EEF5")
    for row_index, row in enumerate(rows, start=1):
        values = list(row)
        while len(values) < len(headers):
            values.append("")
        for col_index, value in enumerate(values[: len(headers)]):
            table.rows[row_index].cells[col_index].text = _text(value)
    doc.add_paragraph("")


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def _ordered_chapters(framework: dict[str, Any]) -> list[dict[str, Any]]:
    chapters = [chapter for chapter in (framework.get("chapters") or []) if isinstance(chapter, dict)]
    return sorted(chapters, key=lambda chapter: int(str(chapter.get("chapter_id") or 0)))


def _format_source_refs(refs: list[Any]) -> str:
    bits = []
    for ref in refs:
        if not isinstance(ref, dict):
            continue
        bits.append(
            " · ".join(
                part
                for part in (
                    _text(ref.get("conversation_id")),
                    _text(ref.get("speaker_role")),
                    _text(ref.get("excerpt_pointer")),
                )
                if part
            )
        )
    return "Sources: " + "; ".join(bit for bit in bits if bit)


def _humanize(key: str) -> str:
    return key.replace("_", " ").strip().title()


def _html_quality(framework: dict[str, Any], labels: dict[str, str]) -> str:
    scores = framework.get("quality_scores") or {}
    rows = "".join(
        f"<tr><th>{html.escape(key.replace('_', ' ').title())}</th>"
        f"<td>{html.escape(str(scores.get(key, '')))}</td></tr>"
        for key in ("opportunity_rating", "conversation_quality", "build_readiness")
        if key in scores
    )
    return f"<h2>{html.escape(labels['quality'])}</h2><table>{rows}</table>"


def _html_body(body: Any) -> list[str]:
    if isinstance(body, str):
        return [f"<p>{html.escape(_text(body))}</p>"] if body.strip() else []
    parts: list[str] = []
    for block in body or []:
        if isinstance(block, str):
            parts.append(f"<p>{html.escape(_text(block))}</p>")
        elif isinstance(block, dict):
            parts.append(_html_block(block))
    return parts


def _html_block(block: dict[str, Any]) -> str:
    kind = str(block.get("block") or "")
    if kind == "prose":
        return f"<p>{html.escape(_text(block.get('text')))}</p>"
    if kind == "bullets":
        items = "".join(f"<li>{html.escape(_text(item))}</li>" for item in block.get("items") or [])
        return f"<ul>{items}</ul>"
    if kind == "callout":
        return f"<p><em>{html.escape(_text(block.get('kind') or 'Note').title())}. {html.escape(_text(block.get('text')))}</em></p>"
    if kind == "process_flow":
        items = "".join(
            f"<li>{html.escape(_text(node.get('label') if isinstance(node, dict) else node))}</li>"
            for node in block.get("nodes") or []
        )
        return f"<ol>{items}</ol>"
    if kind == "ai_split":
        used = "<br>".join(html.escape(_text(item)) for item in block.get("used_for") or [])
        not_used = "<br>".join(html.escape(_text(item)) for item in block.get("not_used_for") or [])
        return (
            "<table><tr><th>AI is used for</th><th>AI is not used for</th></tr>"
            f"<tr><td>{used}</td><td>{not_used}</td></tr></table>"
        )
    if kind in {"kv_rows", "sensitivity"}:
        rows = "".join(
            f"<tr><td>{html.escape(_text(row.get('label')))}</td>"
            f"<td>{html.escape(_text(row.get('value') or row.get('detail')))}</td></tr>"
            for row in block.get("rows") or []
            if isinstance(row, dict)
        )
        return f"<table>{rows}</table>"
    if kind == "table":
        header = "".join(f"<th>{html.escape(_text(col))}</th>" for col in block.get("columns") or [])
        rows = "".join(
            "<tr>" + "".join(f"<td>{html.escape(_text(cell))}</td>" for cell in row) + "</tr>"
            for row in block.get("rows") or []
            if isinstance(row, list)
        )
        return f"<table><tr>{header}</tr>{rows}</table>"
    if kind == "score_bars":
        rows = "".join(
            "<tr>"
            f"<td>{html.escape(_text(item.get('name')))}</td>"
            f"<td>{html.escape(str(item.get('score') or ''))}</td>"
            f"<td>{html.escape(_text(item.get('explanation')))}</td>"
            "</tr>"
            for item in block.get("items") or []
            if isinstance(item, dict)
        )
        return f"<table>{rows}</table>"
    if kind == "timeline":
        rows = "".join(
            f"<tr><td>{html.escape(_text(week.get('id')))}</td>"
            f"<td>{html.escape('; '.join(_text(item) for item in week.get('items') or []))}</td></tr>"
            for week in block.get("weeks") or []
            if isinstance(week, dict)
        )
        return f"<table>{rows}</table>"
    if kind == "glossary":
        rows = "".join(
            f"<tr><td>{html.escape(_text(term.get('term')))}</td>"
            f"<td>{html.escape(_text(term.get('meaning')))}</td></tr>"
            for term in block.get("terms") or []
            if isinstance(term, dict)
        )
        return f"<table>{rows}</table>"
    bits = " ".join(html.escape(_text(value)) for value in block.values() if _text(value))
    return f"<p>{bits}</p>" if bits else ""


def _html_open_items(framework: dict[str, Any], labels: dict[str, str]) -> str:
    items = [item for item in (framework.get("open_items") or []) if isinstance(item, dict)]
    if not items:
        return f"<p>{html.escape(labels['no_open_items'])}</p>"
    rows = "".join(
        "<tr>"
        f"<td>{html.escape(_text(item.get('description')))}</td>"
        f"<td>{html.escape(_text(item.get('item_type')))}</td>"
        f"<td>{html.escape(_text(item.get('owner')))}</td>"
        "</tr>"
        for item in items
    )
    return f"<table>{rows}</table>"
